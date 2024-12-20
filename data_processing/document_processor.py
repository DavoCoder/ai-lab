# Copyright 2024-2025 DavoCoder
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

# document_processor.py
import json
import re
from pathlib import Path
from typing import List, Dict, Any, Tuple
import pandas as pd
import docx
import PyPDF2
from bs4 import BeautifulSoup

class DocumentProcessor:
    SUPPORTED_EXTENSIONS = {
        "PDF": ".pdf",
        "TXT": ".txt",
        "DOCX": ".docx",
        "HTML": ".html",
        "JSON": ".json"
    }

    def clean_text(self, content: str, options: Dict[str, bool]) -> str:
        """Apply text cleaning operations based on options."""
        try:
            if options.get("remove_special_chars"):
                content = re.sub(r'[^a-zA-Z0-9\s]', '', content)
            if options.get("remove_extra_spaces"):
                content = re.sub(r'\s+', ' ', content)
            if options.get("remove_urls"):
                content = re.sub(r'http\S+|www.\S+', '', content)
            if options.get("remove_emails"):
                content = re.sub(r'\S+@\S+', '', content)
            if options.get("normalize_whitespace"):
                content = ' '.join(content.split())
            if options.get("lowercase"):
                content = content.lower()
            return content
        except Exception as e:
            raise Exception(f"Error cleaning text: {str(e)}") from e

    def split_document(self, content: str, split_config: Dict[str, Any]) -> List[str]:
        """Split document based on configuration."""
        try:
            method = split_config["method"]
            chunks = []
            
            if method == "By Paragraphs":
                chunks = [p.strip() for p in content.split("\n\n") if p.strip()]
            elif method == "By Sentences":
                chunks = [s.strip() for s in re.split(r'[.!?]+', content) if s.strip()]
            elif method == "By Word Count":
                words = content.split()
                words_per_chunk = split_config.get("words_per_chunk", 500)
                chunks = [' '.join(words[i:i + words_per_chunk]) 
                         for i in range(0, len(words), words_per_chunk)]
            elif method == "By Custom Delimiter":
                delimiter = split_config.get("delimiter", "\n")
                chunks = [c.strip() for c in content.split(delimiter) if c.strip()]
                
            return chunks
        except Exception as e:
            raise Exception(f"Error splitting document: {str(e)}") from e

    def convert_format(self, content: str, target_format: str) -> Tuple[str, str]:
        """Convert content to target format."""
        try:
            if target_format == "TXT":
                return content, ".txt"
            elif target_format == "JSON":
                return json.dumps({"content": content}), ".json"
            elif target_format == "CSV":
                df = pd.DataFrame({"content": [content]})
                return df.to_csv(index=False), ".csv"
            elif target_format == "HTML":
                return f"<html><body><pre>{content}</pre></body></html>", ".html"
            else:
                raise ValueError(f"Unsupported target format: {target_format}")
        except Exception as e:
            raise Exception(f"Error converting format: {str(e)}") from e

    def extract_metadata(self, file: Any) -> Dict[str, Any]:
        """Extract metadata from file."""
        try:
            metadata = {
                "filename": file.name,
                "size": file.size,
                "type": Path(file.name).suffix.lower()
            }
            
            if metadata["type"] == ".pdf":
                pdf_reader = PyPDF2.PdfReader(file)
                metadata.update({
                    "pages": len(pdf_reader.pages),
                    "info": pdf_reader.metadata
                })
            elif metadata["type"] == ".docx":
                doc = docx.Document(file)
                metadata.update({
                    "paragraphs": len(doc.paragraphs),
                    "sections": len(doc.sections)
                })
                
            return metadata
        except Exception as e:
            raise Exception(f"Error extracting metadata: {str(e)}") from e

    def read_file_content(self, file: Any) -> str:
        """Read content from file."""
        try:
            file_ext = Path(file.name).suffix.lower()
            
            if file_ext == ".txt":
                return file.getvalue().decode()
            elif file_ext == ".pdf":
                pdf_reader = PyPDF2.PdfReader(file)
                return " ".join(page.extract_text() for page in pdf_reader.pages)
            elif file_ext == ".docx":
                doc = docx.Document(file)
                return " ".join(paragraph.text for paragraph in doc.paragraphs)
            elif file_ext == ".html":
                soup = BeautifulSoup(file.getvalue().decode(), 'html.parser')
                return soup.get_text()
            elif file_ext == ".json":
                data = json.loads(file.getvalue().decode())
                return json.dumps(data, indent=2)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        except Exception as e:
            raise Exception(f"Error reading file: {str(e)}") from e

    def batch_process(self, file: Any, operations: List[str]) -> Dict[str, Any]:
        """Process file with multiple operations."""
        try:
            result = {"filename": file.name, "operations": {}}
            
            if "Clean Text" in operations:
                content = self.read_file_content(file)
                cleaned_content = self.clean_text(content, {
                    "remove_special_chars": True,
                    "normalize_whitespace": True
                })
                result["operations"]["clean_text"] = cleaned_content
                
            if "Extract Metadata" in operations:
                metadata = self.extract_metadata(file)
                result["operations"]["metadata"] = metadata
                
            if "Convert Format" in operations:
                content = self.read_file_content(file)
                result["operations"]["converted"] = {
                    "txt": content,
                    "json": json.dumps({"content": content})
                }
                
            return result
        except Exception as e:
            raise Exception(f"Error in batch processing: {str(e)}") from e
